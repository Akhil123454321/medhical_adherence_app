"""
Generate three Word documents from user-instructions.md:
  - admin-instructions.docx
  - patient-instructions.docx
  - chw-instructions.docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
DOCS_DIR = Path("/Users/akhilsk123/Desktop/my_projects/sonak_ui/docs")
MD_FILE  = DOCS_DIR / "user-instructions.md"

COLOR_TITLE    = RGBColor(67,  56, 202)   # indigo
COLOR_H2       = RGBColor(30,  64, 175)   # dark blue
COLOR_H3       = RGBColor(0,   0,   0)    # black
COLOR_BODY     = RGBColor(31,  31,  31)   # near-black
COLOR_NOTE_BG  = RGBColor(200, 200, 200)  # (unused – border only)

FONT_BODY   = "Calibri"
FONT_CODE   = "Courier New"
SIZE_TITLE  = Pt(22)
SIZE_H2     = Pt(14)
SIZE_H3     = Pt(12)
SIZE_BODY   = Pt(11)


# ---------------------------------------------------------------------------
# Helpers — low-level XML
# ---------------------------------------------------------------------------

def set_paragraph_left_border(para, color="AAAAAA", size=12, space=144):
    """Add a left border to a paragraph (for blockquote/note styling)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_horizontal_rule(doc):
    """Insert a thin horizontal rule paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "8B8FA8")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def add_page_numbers(doc):
    """Add a centered page-number field in the footer."""
    section = doc.sections[0]
    footer  = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.clear()

    run = footer_para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    run2 = footer_para.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run2._r.append(instrText)

    run3 = footer_para.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run3._r.append(fldChar2)

    for run in footer_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)


def set_margins(doc, margin_inches=1.0):
    for section in doc.sections:
        section.top_margin    = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin   = Inches(margin_inches)
        section.right_margin  = Inches(margin_inches)


# ---------------------------------------------------------------------------
# Inline-markup parser
# ---------------------------------------------------------------------------

def add_inline_runs(para, text):
    """
    Parse **bold**, `code`, and combinations, then add runs to para.
    Supports mixed inline markup in a single string.
    """
    # Tokenise: split on **...** and `...`
    tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = para.add_run(token[2:-2])
            run.bold = True
            run.font.name  = FONT_BODY
            run.font.size  = SIZE_BODY
            run.font.color.rgb = COLOR_BODY
        elif token.startswith('`') and token.endswith('`'):
            run = para.add_run(token[1:-1])
            run.font.name  = FONT_CODE
            run.font.size  = Pt(10)
            run.font.color.rgb = RGBColor(180, 60, 60)
        else:
            if token:
                run = para.add_run(token)
                run.font.name  = FONT_BODY
                run.font.size  = SIZE_BODY
                run.font.color.rgb = COLOR_BODY


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_doc(title: str, lines: list[str]) -> Document:
    doc = Document()
    set_margins(doc)

    # Default style tweaks
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = SIZE_BODY

    # --- Title ---
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(4)
    run = title_para.add_run(title)
    run.bold = True
    run.font.name  = FONT_BODY
    run.font.size  = SIZE_TITLE
    run.font.color.rgb = COLOR_TITLE

    # Horizontal rule after title
    add_horizontal_rule(doc)

    # Page numbers in footer
    add_page_numbers(doc)

    # --- Parse lines ---
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip blank lines (we handle spacing via paragraph formatting)
        if line.strip() == "":
            i += 1
            continue

        # --- Horizontal rule ---
        if line.strip() in ("---", "***", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # --- H2 (##) ---
        if line.startswith("## "):
            text = line[3:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after  = Pt(4)
            run = para.add_run(text)
            run.bold = True
            run.font.name  = FONT_BODY
            run.font.size  = SIZE_H2
            run.font.color.rgb = COLOR_H2
            i += 1
            continue

        # --- H3 (###) ---
        if line.startswith("### "):
            text = line[4:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after  = Pt(2)
            run = para.add_run(text)
            run.bold = True
            run.font.name  = FONT_BODY
            run.font.size  = SIZE_H3
            run.font.color.rgb = COLOR_H3
            i += 1
            continue

        # --- Numbered list (1. 2. etc.) ---
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            text = m.group(2).strip()
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.left_indent   = Inches(0.25)
            para.paragraph_format.space_before  = Pt(2)
            para.paragraph_format.space_after   = Pt(2)
            add_inline_runs(para, text)
            i += 1
            continue

        # --- Bullet list (- or *) ---
        m = re.match(r'^[-*]\s+(.*)', line)
        if m:
            text = m.group(1).strip()
            # Check for sub-bullet (indented)
            indent_level = 0
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent  = Inches(0.25 + 0.25 * indent_level)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            add_inline_runs(para, text)
            i += 1
            continue

        # --- Indented bullet (  - sub item) ---
        m = re.match(r'^\s{2,}[-*]\s+(.*)', line)
        if m:
            text = m.group(1).strip()
            para = doc.add_paragraph(style="List Bullet 2")
            para.paragraph_format.left_indent  = Inches(0.5)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            add_inline_runs(para, text)
            i += 1
            continue

        # --- Blockquote / Note (>) ---
        if line.startswith("> "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent  = Inches(0.3)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after  = Pt(6)
            set_paragraph_left_border(para, color="AAAAAA", size=18, space=144)
            run = para.add_run(text)
            run.italic = True
            run.font.name  = FONT_BODY
            run.font.size  = SIZE_BODY
            run.font.color.rgb = RGBColor(80, 80, 80)
            # Re-parse bold inside blockquotes
            para.clear()
            # parse inline in italic context
            tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
            for token in tokens:
                if token.startswith('**') and token.endswith('**'):
                    r = para.add_run(token[2:-2])
                    r.bold   = True
                    r.italic = True
                    r.font.name  = FONT_BODY
                    r.font.size  = SIZE_BODY
                    r.font.color.rgb = RGBColor(60, 60, 60)
                elif token.startswith('`') and token.endswith('`'):
                    r = para.add_run(token[1:-1])
                    r.font.name  = FONT_CODE
                    r.font.size  = Pt(10)
                    r.font.color.rgb = RGBColor(150, 50, 50)
                elif token:
                    r = para.add_run(token)
                    r.italic = True
                    r.font.name  = FONT_BODY
                    r.font.size  = SIZE_BODY
                    r.font.color.rgb = RGBColor(70, 70, 70)
            set_paragraph_left_border(para, color="AAAAAA", size=18, space=144)
            i += 1
            continue

        # --- Italic footer line (*...*) ---
        if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            text = line.strip("*").strip()
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            run = para.add_run(text)
            run.italic = True
            run.font.name  = FONT_BODY
            run.font.size  = Pt(10)
            run.font.color.rgb = RGBColor(120, 120, 120)
            i += 1
            continue

        # --- Regular paragraph ---
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after  = Pt(4)
        add_inline_runs(para, line.strip())
        i += 1

    return doc


# ---------------------------------------------------------------------------
# Section extractor
# ---------------------------------------------------------------------------

def extract_section(md_text: str, heading: str) -> list[str]:
    """
    Return lines between '# <heading>' and the next top-level '#' heading.
    Skips the heading line itself (we supply our own title).
    """
    lines   = md_text.splitlines()
    in_sec  = False
    result  = []
    pattern = re.compile(r'^#\s+' + re.escape(heading) + r'\s*$')
    stop    = re.compile(r'^#\s+\w')   # another top-level heading

    for line in lines:
        if pattern.match(line):
            in_sec = True
            continue
        if in_sec:
            if stop.match(line) and not pattern.match(line):
                break
            result.append(line)

    return result


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    md_text = MD_FILE.read_text(encoding="utf-8")

    sections = {
        "admin-instructions.docx": {
            "heading":  "Admin Instructions",
            "title":    "MedAdhere — Admin Instructions",
        },
        "patient-instructions.docx": {
            "heading":  "Patient Instructions",
            "title":    "MedAdhere — Patient Instructions",
        },
        "chw-instructions.docx": {
            "heading":  "CHW Instructions",
            "title":    "MedAdhere — CHW Instructions",
        },
    }

    for filename, cfg in sections.items():
        lines = extract_section(md_text, cfg["heading"])
        doc   = build_doc(cfg["title"], lines)
        out   = DOCS_DIR / filename
        doc.save(str(out))
        print(f"Saved: {out}  ({len(lines)} source lines)")

    print("Done.")


if __name__ == "__main__":
    main()
